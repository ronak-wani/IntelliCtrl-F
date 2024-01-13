import magic
from flask import Flask, redirect, render_template, request, jsonify
import io
import json
from docx.api import Document
import pytesseract
import os, pyttsx3, speech_recognition
import PIL
from pdfminer.high_level import extract_pages, extract_text
from PIL.Image import Image
import re, PyPDF2
import tempfile
from tempfile import gettempdir
from flask import render_template, redirect, request, app, flash, send_file
from werkzeug.utils import secure_filename
from nlp import ctrl_f

app = Flask(__name__)


@app.route("/", methods=["GET", "POST"])
def index():
        return render_template("index.html")

def userText():
    if request.method == "GET":
        text = request.args.get("code")
        language = request.args.get("language")
        print(text, language)

def get_file_type(filepath):
    mime_type = magic.from_file(filepath, mime=True)
    return mime_type

@app.route("/detect_file_type", methods=["GET", "POST"])
def detect_file_type():
    print(request)
    file=None
    query=None
    if request.method == 'POST':
        query = request.form.get('query')
        if 'fileInput' not in request.files:
            return jsonify({'error': 'No file part'})
        file = request.files['fileInput']
    filename = secure_filename(file.filename)
    filepath = os.path.join(tempfile.gettempdir(), filename)
    file.save(filepath)
    text = None
    if get_file_type(filepath) == 'application/pdf':
        text = pdf_recognition(filepath)
        # return render_template("nlp.html")
    elif get_file_type(filepath) == 'text/plain':
        text = txt_recognition(filepath)
        # return render_template("nlp.html")
    elif get_file_type(filepath) == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text = docx_recognition(filepath)
    elif get_file_type(filepath) == 'image/png' or 'image/jpeg':
        text = image_recognition(filepath)
    else:
        return jsonify({'error': 'Wrong file type'})
    text = text.decode('utf-8').replace('\u200C','').strip()
    sentences = ctrl_f(text, query)
    return jsonify({'sentences': sentences})
@app.route("/text_extract", methods=['GET', 'POST'])
def text_extract():
    if request.method == "GET":
        text = request.args.get("code")
        print(text)
        return render_template("nlp.html")
def pdf_recognition(filepath): #TODO fix
    res_text = ""
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for i in range(len(reader.pages)):
            page = reader.pages[i]
            text = page.extract_text()
            res_text += text + " "
    return res_text
def txt_recognition(filepath):
    with open(filepath, "rb") as f:
        text = f.read()
        return text

def docx_recognition(filepath):
    document = Document(filepath)
    all_text = ""
    for(table) in document.tables:
        for(row) in table.rows:
            print("|".join(cell.text for cell in row.cells))
    for p in document.paragraphs:
        all_text += p.text
        all_text += "\n"
    return all_text

def image_recognition(image: Image):
    myconfig = r"--psm 6 --oem 3"
    text = pytesseract.image_to_string(PIL.Image.open(image), config=myconfig)
    return  text

def text_to_speech(self):
    if request.method == "GET":
        test = request.args.get("tts")
        print(test)
        text_speech = pyttsx3.init()
        text_speech.say(f"{test}")
        text_speech.runAndWait()
        return None

def speech_recognition(self):
    if request.method == "GET":
        recognizer = speech_recognition.Recognizer()
        while True:
                 try:
                     with speech_recognition.Microphone() as mic:
                         recognizer.adjust_for_ambient_noise(mic, duration=0.5)
                         audio = recognizer.listen(mic)
                         text = recognizer.recognize_google(audio)
                         text = text.lower()
                         print(f"Recognized {text}")
                         return json.dumps(text)
                 except speech_recognition.UnknownValueError:
                    # Reset the recognizer
                    recognizer = speech_recognition.Recognizer()
                    # Continue listening
                    continue
                 #Reinitialize the text variable
                 text = ""

if __name__ == "__main__":
    app.run(port=8001)